from io import BytesIO
import streamlit as st
import os
from rag_pipeline import search_llm 
from external_search import  get_final_result
from reportlab.lib.pagesizes import letter
import pickle 
from pathlib import Path
import streamlit_authenticator as stauth
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import RGBColor, Pt
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_exa import ExaSearchRetriever, TextContentsOptions
from langchain_google_genai import ChatGoogleGenerativeAI
import os
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_community.retrievers import TavilySearchAPIRetriever
os.environ["GOOGLE_API_KEY"] = "AIzaSyDhpqwIfCULLr6Gv3s955rxIRXyRtDDyhk"
os.environ["HUGGINGFACEHUB_API_TOKEN"] = "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"
os.environ['TAVILY_API_KEY'] = "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"
st.set_page_config(page_title="💬PPP Benchmark studies Chatbot")


# --- USER AUTHENTIFICATION ---

names = ["Houyem", "Ilef"]
usernames = ["houyem", "test"]


# load hashed passwords
file_path = Path(__file__).parent / "hashed_pw.pkl"
with file_path.open("rb") as file:
    hashed_passwords = pickle.load(file)

authenticator = stauth.Authenticate(names, usernames, hashed_passwords,
    "jade_chatbot", "abcdef", cookie_expiry_days=30)

name, authentication_status, username = authenticator.login("Login", "main")

if authentication_status == False:
    st.error("Username/password is incorrect")

if authentication_status == None:
    st.warning("Please enter your username and password")

if authentication_status:
    # Define the directory where PDFs will be saved
    PDF_DIR = "D:/HenloIlef/2K24/chatbotJade/chatbotJade/allLocal/"
    # Function to save response to PDF
    

    # Initialize session state
    if "questions_button_clicked" not in st.session_state:
        st.session_state.questions_button_clicked = False

    if "conduct_button_clicked" not in st.session_state:
        st.session_state.conduct_button_clicked = False

    with st.sidebar:
        st.sidebar.title(f"Welcome {name}")
        st.title('💬PPP Benchmark studies Chatbot')

        # Add two buttons to the sidebar
        if st.button("Questions about existing documents"):
            # Mark the button as clicked
            st.session_state.questions_button_clicked = True
            # Add a message to the chat history
            st.session_state.messages.append({"role": "assistant", "content": "What questions do you have about existing documents?"})

        if st.button("Conduct new benchmark studies"):
            # Mark the button as clicked
            st.session_state.conduct_button_clicked = True
            # Add a message to the chat history
            st.session_state.messages.append({"role": "assistant", "content": "Let's conduct new benchmark studies. What do you want to know?"})
        authenticator.logout("Logout", "sidebar")
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Display chat messages from history on app rerun
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # React to user input
    prompt = st.chat_input("How can I help you?")
    if prompt:
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(prompt)
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})

        # Call the appropriate function based on the context
        if st.session_state.questions_button_clicked:
            response = search_llm(prompt)
        elif st.session_state.conduct_button_clicked:
            response = get_final_result(prompt)
            


        # Display the response in the chat
        with st.chat_message("assistant"):
            st.markdown(response)
        
               
        # Add the assistant's response to chat history
        st.session_state.messages.append({"role": "assistant", "content": response})
        blue_titles = [
    "1. Présentation du projet",
    "2. Structure contractuelle du projet",
    "3. Leçons tirées/Recommandations/Meilleures pratiques/Erreurs à éviter",
    "Leçons tirées:",
    "3. Leçons tirées",
    "Project presentation",
    "Project presentation:",
    "1. Project presentation",
    "1. Project Presentation",
    "Project presentation",
    "Project presentation:",
    "Project Presentation:",
    "1. Project Overview",
    "Project Contractual Structure",
    "Project Contractual Structure:",
    "2. Project Contractual Structure",
    "Project Contractual Structure:",
    
    "3. Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid",
    "Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid",
    "Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid:",
    "Project Overview",
    
    ]

        def create_docx_with_colored_titles(content, blue_titles):
            doc = Document()
        
            # Traiter chaque ligne du contenu
            for line in content.split("\n"):
            # Nettoyer les "**" au début et à la fin si présents
                clean_line = line.strip("* ")
            
            # Ajouter la ligne au document
                if clean_line in blue_titles:
                # Si la ligne est un des titres spécifiques, l'ajouter en bleu
                    para = doc.add_paragraph()
                    run = para.add_run(clean_line)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Bleu
                    run.font.size = Pt(14)  # Taille de la police
                else:
                # Sinon, ajouter la ligne comme paragraphe normal
                    doc.add_paragraph(clean_line)
        
        # Sauvegarder le document dans un buffer en mémoire
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
        
            return buffer

        # Appeler la fonction et créer le document
        buffer = create_docx_with_colored_titles(response, blue_titles)

        # Créer un bouton de téléchargement pour le document DOCX dans Streamlit
        st.download_button(
            label="Download",
            data=buffer,
            file_name="bechmark_study.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )



mistralF = CTransformers(model='katkout2313/mistral-for-becnhmarks-finetunedFinal',model_file="mistral-for-becnhmarks-finetunedFinal-unsloth.Q4_K_M.gguf",
                            model_type='llama',
                            config={'max_new_tokens': 2048,
                                    'temperature' :0.2,
                                    'top_k':60,
                                    'top_p':0.95,
    }
   