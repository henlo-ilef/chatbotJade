from langchain.callbacks.manager import collect_runs
import docx
import streamlit as st
import os
from rag_pipeline import search_llm 
from external_search import get_response , titre , return_image
#from rag_pipeline import search_llm
#from external_search import titre , return_image , titre_response
from streamlit.web import cli as stcli
import pickle 
from pathlib import Path
from PIL import Image

import streamlit_authenticator as stauth
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from io import BytesIO
import firebase_admin
from firebase_admin import credentials

from firebase_admin import firestore
from account import login, sign
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from streamlit_option_menu import option_menu

from docx import Document
from docx.shared import RGBColor, Pt , Inches
from streamlit_feedback import streamlit_feedback
from langchain.callbacks.tracers.run_collector import RunCollectorCallbackHandler
from langchain.memory import StreamlitChatMessageHistory, ConversationBufferMemory
from langchain.schema.runnable import RunnableConfig
from langsmith import Client
from dotenv import load_dotenv
load_dotenv()
st.set_page_config(page_title="💬PPP Benchmark studies Chatbot")
service_account_path = './jade.json'
os.environ["LANGCHAIN_PROJECT"] = "JADE-AI"
os.environ["LANGCHAIN_API_KEY"] ="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" 
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
client = Client(api_url= "https://api.smith.langchain.com", api_key="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")

def initialize_firebase():
    # Check if any apps are already initialized in this environment
    if not firebase_admin._apps:
        # If no apps, initialize app with credentials
        cred = credentials.Certificate(service_account_path)
        firebase_app = firebase_admin.initialize_app(cred)
    else:
        # If already initialized, use the existing app
        firebase_app = firebase_admin.get_app()
    return firebase_app

# Initialize Firebase
firebase_app = initialize_firebase()
db = firestore.client(app=firebase_app)
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False


if not st.session_state.logged_in:
    selected_form = option_menu(
    menu_title=None,  # Pas de titre pour le menu
    options=["Login", "Sign Up"],  # Les options disponibles
    icons=["box-arrow-in-right", "pencil-square"],  # Icônes pour chaque option
    orientation="horizontal",
    default_index=0 if st.session_state.get('form_type', 'login') == 'login' else 1,
)

# Définir form_type en fonction de la sélection
    st.session_state.form_type = 'login' if selected_form == "Login" else 'signup'
    # Gestion de l'affichage des formulaires
    if st.session_state.form_type == 'signup':
        st.subheader("Sign Up")
        if sign():
             #st.session_state.logged_in = True
            username = st.session_state.get("username", "")
            #st.success("Sign up successful! You are now logged in.")
            #st.rerun()  # Recharger la page pour refléter l'état connecté

    else:
        st.subheader("Login")
        if login():
            st.session_state.logged_in = True
            username = st.session_state.get("username", "")
            st.rerun()  # Recharger la page pour refléter l'état connecté
else : 
    st.empty()

if 'username' in st.session_state:
    username = st.session_state['username']
    user_id = username

# Initialize Firebase

if st.session_state.logged_in: 
    
    # --- USER AUTHENTIFICATION ---


    from google.cloud.firestore import ArrayUnion
    import datetime
    blue_titles = [
    "1. Présentation du projet",
    "1. Présentation du projet:",
    "2. Structure contractuelle du projet",
    "2. Structure contractuelle du projet:",
    "3. Leçons tirées/Recommandations/Meilleures pratiques/Erreurs à éviter",
    "Leçons tirées:",
    "3. Leçons tirées",
    "Project presentation",
    "Project presentation:",
    "1. Project presentation",
    "1. Project Presentation",
    "Project presentation",
    "Project presentation:",
    "Project Presentation:",
    "1. Project Overview",
    "1. Project Overview:",
    "Project Contractual Structure",
    "Project Contractual Structure:",
    "2. Project Contractual Structure",
    "2. Project Contractual Structure:",
    "Project Contractual Structure:",
    
    "3. Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid",
    "Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid",
    "Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid:",
    "3. Lessons Learned/Recommendations/Best Practices/Mistakes to Avoid:",
    "Project Overview",
    
    ]
    def image_to_jpg(image_stream):
    # Convert image bytes to JPEG format
        with Image.open(image_stream) as img:
            img = img.convert("RGB")
            jpg_image_stream = BytesIO()
            img.save(jpg_image_stream, format="JPEG")
            jpg_image_stream.seek(0)
        return jpg_image_stream


    def create_docx_with_colored_titles(content, blue_titles, images):
        doc = Document()
        overview_titles = [
            "1. Project Overview", "Project Overview", "1. Project overview", "project overview",
            "1. Project overview:", "1. Project Overview:", "1. Présentation du projet",
            "1. Présentation du projet:", "1. présentation du projet:"
        ]

        image_index = 0  # Index to track the current image to be placed

        # Process each line of the content
        for line in content.split("\n"):

            # Remove leading and trailing whitespace and asterisks
            clean_line = line.strip().strip("*")  # Stripping '*' here
            clean_line = clean_line.replace("**", "")  # Additional step to remove '**' anywhere in the text

            # Add the line to the document
            para = doc.add_paragraph()
            run = para.add_run(clean_line)
            if clean_line in blue_titles:
                run.font.color.rgb = RGBColor(0, 0, 255)  # Apply blue color to titles
                run.font.size = Pt(14)  # Set font size for titles

            # Check if the current line is an overview title and if an image should be added
            if clean_line.startswith(tuple(overview_titles)) and image_index < len(images):
                image_url = images[image_index]
                response = requests.get(image_url)
                if response.status_code == 200:
                    image_stream = BytesIO(response.content)  # Create a stream for the image
                    # Convert image to JPEG if needed
                    image_stream = image_to_jpg(image_stream)
                    doc.add_picture(image_stream, width=Inches(5.5))  # Add picture to the document
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the image
                    image_index += 1  # Move to the next image

        # Save the document to a memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer
        
    def save_message(user_id, role, content, new_chat=False, chat_id=None):
        chats_ref = db.collection('users').document(user_id).collection('chats')
        if new_chat:
                # Start a new chat document
            chat_ref = chats_ref.document()
            chat_ref.set({
                'started_at': firestore.SERVER_TIMESTAMP,
                'updated_at': firestore.SERVER_TIMESTAMP,
                'status': 'active',
                'messages': [{
                    'timestamp': datetime.datetime.now(),
                    #'sender': role,
                     'role':role,
                    'content': content,
                    'type': 'text'
                }]
            })
        else:
                # Use the existing chat_id if provided
            if chat_id:
                chat_ref = chats_ref.document(chat_id)
            else:
                    # Find the most recent chat document
                chat_doc = chats_ref.order_by('updated_at', direction=firestore.Query.DESCENDING).limit(1).get()
                if chat_doc and list(chat_doc):
                    chat_doc = chat_doc[0]
                    chat_ref = chats_ref.document(chat_doc.id)

            new_message = {
                'timestamp': datetime.datetime.now(),
                #'sender': role,
                'content': content,
                'type': 'text',
                'role': role
            }
            chat_ref.update({
                'messages': ArrayUnion([new_message]),
                'updated_at': firestore.SERVER_TIMESTAMP
            })





    def get_history(user_id):
        chats_ref = db.collection('users').document(user_id).collection('chats')
        all_chats = []
        chats = chats_ref.stream()  # Récupère tous les documents de chat pour cet utilisateur

        for chat in chats:
            chat_data = chat.to_dict()
            messages = chat_data.get('messages', [])
            all_chats.append({
                'chat_id': chat.id,  # Inclure l'ID du document de chat
                'messages': messages
            })

        return all_chats
    
    def get_chat_title(chat_id, message_content):
        if chat_id not in st.session_state.chat_titles:
            # Supposons que `titre()` est votre fonction qui génère un titre
            st.session_state.chat_titles[chat_id] = titre(message_content)
        return st.session_state.chat_titles[chat_id]

    if 'chat_titles' not in st.session_state:
        st.session_state.chat_titles = {}

    if "questions_button_clicked" not in st.session_state:
        st.session_state.questions_button_clicked = False

    if "conduct_button_clicked" not in st.session_state:
        st.session_state.conduct_button_clicked = False



        #user_id = username  # Example to link chats to user id
    user_id=username
    with st.sidebar:
        st.title(f"Welcome {username}")
        if st.button("Questions about existing documents"):
                    # Mark the button as clicked
            st.session_state.questions_button_clicked = True
            st.session_state.messages = []
                    # Add a message to the chat history
            st.session_state.messages.append({"role": "assistant", "content": "What questions do you have about existing documents?"})

        if st.button("Conduct new benchmark studies"):
                    # Mark the button as clicked
            st.session_state.conduct_button_clicked = True
            st.session_state.start_new_chat = True
            st.session_state.selected_chat = None 
            st.session_state.messages = []
            # Ajouter un message initial pour le chat
            initial_message = "Let's conduct new benchmark studies. What do you want to know?"
            st.session_state.messages.append({"role": "assistant", "content": initial_message})
            # Sauvegarder le message dans un nouveau chat
            save_message(user_id, "assistant", initial_message, new_chat=True)
        st.title("Chat History")
        all_chats = get_history(user_id)
        if all_chats:
            for chat in all_chats:
                   
                if chat['messages']: 
                    if len(chat['messages']) > 1:
                        #msg = chat['messages'][1]['content']
                        #titre_result=titre(msg)
                        titre_result=get_chat_title(chat['chat_id'], chat['messages'][1]['content'])
                    else:
                        titre_result=chat['messages'][0]['content']

                       
                    if st.button(f"{titre_result}", key=chat['chat_id']):
                        st.session_state.conduct_button_clicked = False  # Assurez-vous de désactiver l'état de la nouvelle conversation
                        st.session_state.selected_chat = chat  # Mettre à jour le chat sélectionné
                        st.session_state.messages = chat['messages']  # Mettre à jour les messages à afficher
                        st.session_state.start_new_chat = False
                        st.rerun()
                        
                       
                        
                else:
                    st.write(f"Chat {chat['chat_id']} has no messages yet.")
        else:
            st.write("No chat history found.")
                
    if 'start_new_chat' not in st.session_state:
        st.session_state.start_new_chat = False

    if 'selected_chat' in st.session_state and not st.session_state.start_new_chat:
        chat = st.session_state.selected_chat
        
        st.write(f"Chat started on: {chat['messages'][0]['timestamp']}")
        for message in chat['messages']:
            with st.chat_message(message["role"]):
                st.markdown(message['content'])

    if "messages" not in st.session_state:
        st.session_state.messages = []
            # Display chat messages from history on app rerun
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    prompt = st.chat_input("How can I help you?")
    if prompt:
        with st.chat_message("user"):
            st.markdown(prompt)
            if 'selected_chat_id' in st.session_state:
                save_message(user_id, "user", prompt, chat_id=st.session_state.selected_chat_id)
            else :
                save_message(user_id, "user", prompt)

                    
                
        st.session_state.messages.append({"role": "user", "content": prompt})
            
        
                
        if st.session_state.conduct_button_clicked:
            
            response ,image , projects= get_response(prompt)
            if 'selected_chat_id' in st.session_state:
                save_message(user_id, "assistant", response, chat_id=st.session_state.selected_chat_id)
            else:
                save_message(user_id, "assistant", response)
                
            with st.chat_message("assistant"):
                st.markdown(response)
            st.session_state.messages.append({"role": "assistant", "content": response})
            
            #t=titre_response(response)
            images=[]
            for project in projects:
                image1 = return_image(project)
                images.append(image1)

        # Appeler la fonction et créer le document
            buffer = create_docx_with_colored_titles(response, blue_titles, images)

        # Créer un bouton de téléchargement pour le document DOCX dans Streamlit
            st.download_button(
                label="Download",
                data=buffer,
                file_name="bechmark_study.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            

    
        elif st.session_state.questions_button_clicked:
            with st.chat_message("assistant"):
                chat_history = []
                rag_chain = search_llm()
                response = rag_chain.invoke({"input": prompt, "chat_history": chat_history})
                from langchain_core.messages import HumanMessage
                chat_history.extend([HumanMessage(content=prompt), response["answer"]])
                st.markdown(prompt)
                
                
                docs = []
                for document in response["context"]:
                    docs.append(document)
                response_final = response["answer"] 
                st.markdown(response_final)
            st.session_state.messages.append({"role": "assistant", "content": response["answer"]})


            
